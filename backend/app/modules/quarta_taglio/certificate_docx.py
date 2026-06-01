from __future__ import annotations

import zipfile
import zlib
import re
import tempfile
from collections.abc import Sequence
from pathlib import Path

import fitz
from PIL import Image
from docxcompose.composer import Composer
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

from app.core.users.models import User
from app.modules.quarta_taglio.schemas import QuartaTaglioDetailResponse


APP_ROOT = Path(__file__).resolve().parents[2]
ASSET_ROOT = APP_ROOT / "assets" / "certificates"
LOGO_PATH = ASSET_ROOT / "forgialluminio_logo.png"
QUALITY_MANAGER_SIGNATURE_PATH = ASSET_ROOT / "quality_manager_signature.png"
HEADER_WIDTH = Inches(7.1)
HEADER_LOGO_WIDTH = Inches(2.18)
HEADER_LOGO_WIDTH_EMU = 1993392
HEADER_LOGO_HEIGHT_EMU = 927741
HEADER_FLOW_COLUMN_WIDTHS = [Inches(2.05), Inches(2.525), Inches(2.525)]
HEADER_FLOW_ROW_HEIGHT = Twips(395)
HEADER_FLOW_LINE_SPACING = 220 / 240
CHEMISTRY_TABLE_WIDTH_INCHES = 7.1
CHEMISTRY_LABEL_COLUMN_WIDTH_INCHES = 0.70

PROPERTY_HEADER_LABELS = {
    "HB": "HB",
    "Rp0.2": "Rp 0,20\n(N/mm²)",
    "Rp0.20": "Rp 0,20\n(N/mm²)",
    "Rm": "Rm\n(N/mm²)",
    "A%": "A(%)",
    "IACS%": "IACS(%)",
    "Rp0.2 / Rm": "Rp/Rm",
    "Rp0.2/Rm": "Rp/Rm",
    "diametro": "Ø(mm)",
    "Diametro": "Ø(mm)",
    "Ø": "Ø(mm)",
    "S": "S(mm²)",
}
PROPERTY_WORD_FIELDS = ["HB", "diametro", "S", "Rp0.2", "Rm", "A%", "Rp0.2 / Rm"]
CONTENT_CONTROL_TAGS = (
    "CERT_NUMBER",
    "CERT_DATE",
    "SUPPLIER_REF",
    "PURCHASER",
    "ORDER_CLIENT",
    "CONFIRM_ORDER",
    "COD_F3_RAW",
    "RAW_DESCRIPTION",
    "DDT_RAW",
    "QUANTITY_RAW",
    "COD_F3_FINISHED",
    "FINISHED_DESCRIPTION",
    "DDT_FINISHED",
    "QUANTITY_FINISHED",
)


def build_additional_page_template_docx(output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    _set_document_sections_layout(document)
    _clear_header_footer(document.sections[0].header)
    _clear_header_footer(document.sections[0].footer)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title_run = title.add_run("Titolo pagina aggiuntiva")
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(12)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(8)
    subtitle_run = subtitle.add_run("Testo o descrizione breve")
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = RGBColor(90, 90, 90)

    data_table = document.add_table(rows=6, cols=4)
    data_table.style = "Table Grid"
    data_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_width(data_table, Inches(7.1))
    _set_table_cell_margins(data_table, left=45, right=45, top=25, bottom=25)
    headers = ("Campo", "Valore", "Campo", "Valore")
    for index, label in enumerate(headers):
        paragraph = data_table.rows[0].cells[index].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(label)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(9)
    for row in data_table.rows:
        row.height = Inches(0.23)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    _set_table_font(data_table, size=9)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(8)
    spacer.paragraph_format.space_after = Pt(4)

    image_table = document.add_table(rows=1, cols=1)
    image_table.style = "Table Grid"
    image_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_width(image_table, Inches(7.1))
    _set_table_cell_margins(image_table, left=45, right=45, top=35, bottom=35)
    image_row = image_table.rows[0]
    image_row.height = Inches(3.35)
    image_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    image_cell = image_row.cells[0]
    image_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    image_paragraph = image_cell.paragraphs[0]
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_run = image_paragraph.add_run("Area immagine / tabella / allegato")
    image_run.font.name = "Arial"
    image_run.font.size = Pt(10)
    image_run.font.color.rgb = RGBColor(120, 120, 120)

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(8)
    note.paragraph_format.space_after = Pt(0)
    note_run = note.add_run("Contenuto modificabile: usa questo modello per mantenere larghezze e spazi compatibili con il certificato.")
    note_run.font.name = "Arial"
    note_run.font.size = Pt(8)
    note_run.font.color.rgb = RGBColor(120, 120, 120)

    document.save(output_path)


def build_forgialluminio_draft_docx(
    *,
    detail: QuartaTaglioDetailResponse,
    output_path: Path,
    draft_number: str,
    certified_by: User,
    quality_manager: User | None,
    additional_pages_path: Path | None = None,
    pdf_attachments: Sequence[tuple[Path, str]] | None = None,
) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    _set_document_sections_layout(document)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(9)

    _apply_document_shell(document, detail=detail, draft_number=draft_number, certified_by=certified_by, quality_manager=quality_manager)
    _add_chemistry_table(document, detail=detail)
    _add_notes(document, detail=detail)
    _add_properties_table(document, detail=detail)

    document.save(output_path)
    if additional_pages_path is not None and additional_pages_path.exists():
        _append_docx_body(
            output_path,
            additional_pages_path,
            detail=detail,
            draft_number=draft_number,
            certified_by=certified_by,
            quality_manager=quality_manager,
        )
    if pdf_attachments:
        _append_pdf_attachments(
            output_path,
            pdf_attachments,
            certified_by=certified_by,
            quality_manager=quality_manager,
        )
    normalize_certificate_docx_layout(output_path)


def _set_document_sections_layout(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Inches(2.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)
        section.header_distance = Inches(0.48)
        section.footer_distance = Inches(0.12)
        section.different_first_page_header_footer = False
        _remove_page_number_restart(section)


def _apply_document_shell(
    document: Document,
    *,
    detail: QuartaTaglioDetailResponse,
    draft_number: str,
    certified_by: User,
    quality_manager: User | None,
) -> None:
    _set_document_sections_layout(document)
    for index, section in enumerate(document.sections):
        if index == 0:
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False
            _fill_document_header(section.header, detail=detail, draft_number=draft_number)
            _fill_signature_footer(section.footer, certified_by=certified_by, quality_manager=quality_manager, include_operator=True)
        else:
            section.header.is_linked_to_previous = True
            section.footer.is_linked_to_previous = True


def _fill_document_header(header, *, detail: QuartaTaglioDetailResponse, draft_number: str) -> None:
    certificate_header = detail.header or {}
    _clear_header_footer(header)
    table = header.add_table(rows=1, cols=3, width=HEADER_WIDTH)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _clear_table_borders(table)
    _set_column_widths(table, [Inches(1.2), Inches(4.4), Inches(1.5)])
    left, center, right = table.rows[0].cells
    if LOGO_PATH.exists():
        paragraph = center.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run().add_picture(str(LOGO_PATH), width=HEADER_LOGO_WIDTH)
    else:
        paragraph = center.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run("Forgialluminio3 s.r.l.").bold = True
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Mod. 065 Rev. 02")
    run.font.size = Pt(5.5)
    p = right.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Pag. ")
    run.font.size = Pt(9)
    _add_field(p, "PAGE")
    run = p.add_run(" / ")
    run.font.size = Pt(9)
    _add_field(p, "NUMPAGES")

    title_table = header.add_table(rows=1, cols=2, width=HEADER_WIDTH)
    _clear_table_borders(title_table)
    title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_width(title_table, HEADER_WIDTH)
    _set_column_widths(title_table, [Inches(3.15), Inches(3.95)])
    left_title, right_title = title_table.rows[0].cells
    left_paragraph = left_title.paragraphs[0]
    left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_paragraph.paragraph_format.space_before = Pt(0)
    left_paragraph.paragraph_format.space_after = Pt(0)
    left_paragraph.paragraph_format.line_spacing = 0.85
    left_run = left_paragraph.add_run("EN 10204 - 3.1")
    left_run.bold = True
    left_run.font.size = Pt(13)
    _set_cell_no_wrap(left_title)

    right_paragraph = right_title.paragraphs[0]
    right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_paragraph.paragraph_format.space_before = Pt(0)
    right_paragraph.paragraph_format.space_after = Pt(0)
    right_paragraph.paragraph_format.line_spacing = 0.85
    prefix_run = right_paragraph.add_run("Certificate n°: ")
    prefix_run.font.size = Pt(10)
    _add_content_control_run(
        right_paragraph,
        tag="CERT_NUMBER",
        text=draft_number,
        font_name="Arial",
        size=10,
    )
    dated_run = right_paragraph.add_run(" dated ")
    dated_run.font.size = Pt(10)
    _add_content_control_run(
        right_paragraph,
        tag="CERT_DATE",
        text=certificate_header.get("data_certificato") or "",
        font_name="Arial",
        size=10,
    )
    _set_cell_no_wrap(right_title)

    _add_certificate_header_flow_table(header, certificate_header)


def _add_chemistry_intro(document: Document, *, detail: QuartaTaglioDetailResponse) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(16)
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run("•  Chemical composition: acc. to EN UNI 573-3")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

    info_table = document.add_table(rows=1, cols=3)
    _clear_table_borders(info_table)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_width(info_table, Inches(6.7))
    _set_column_widths(info_table, [Inches(2.3), Inches(2.5), Inches(1.9)])
    alloy_cell, charge_cell, ref_cell = info_table.rows[0].cells
    _add_inline_label_value(alloy_cell.paragraphs[0], "Alloy:", _alloy_label(detail))
    _add_inline_label_value(charge_cell.paragraphs[0], "Charge:", detail.cod_odp)
    _add_inline_label_value(
        ref_cell.paragraphs[0],
        "Ref.:",
        (detail.header or {}).get("supplier_ref") or "-",
        control_tag="SUPPLIER_REF",
    )
    for cell in (alloy_cell, charge_cell, ref_cell):
        _set_cell_no_wrap(cell)
    document.add_paragraph().paragraph_format.space_after = Pt(4)


def _add_materials_table(document: Document, *, detail: QuartaTaglioDetailResponse) -> None:
    document.add_page_break()
    _add_section_title(document, "Supplier certificates / Certificati fornitore")
    table = _new_table(document, ["CDQ", "Colata", "Cod. art.", "Quantita", "Lotti"])
    for item in detail.materials:
        row = table.add_row().cells
        row[0].text = item.cdq or "-"
        row[1].text = item.colata or "-"
        row[2].text = item.cod_art or "-"
        row[3].text = _format_quantity(item.qta_totale)
        row[4].text = ", ".join(item.cod_lotti or []) or "-"


def _add_chemistry_table(document: Document, *, detail: QuartaTaglioDetailResponse) -> None:
    _add_chemistry_intro(document, detail=detail)
    chemistry = [item for item in detail.chemistry if item.status != "not_in_standard"]
    if not chemistry:
        document.add_paragraph("No chemical elements available from selected standard.")
        return
    table = _new_table(document, [""] + [item.field for item in chemistry])
    min_row = table.add_row().cells
    max_row = table.add_row().cells
    value_row = table.add_row().cells
    min_row[0].text = "%\u00a0min"
    max_row[0].text = "%\u00a0max"
    value_row[0].text = "%\u00a0val"
    for index, item in enumerate(chemistry, start=1):
        min_row[index].text = _format_chemistry_value(item.standard_min)
        max_row[index].text = _format_chemistry_value(item.standard_max)
        value_row[index].text = _format_chemistry_value(item.value)
    _set_table_font(table, size=10)
    _set_table_cell_margins(table, left=25, right=25)
    for row in table.rows:
        for cell in row.cells:
            _set_cell_no_wrap(cell)
    _bold_row(value_row)
    _set_table_width_with_first_column(
        table,
        total_inches=CHEMISTRY_TABLE_WIDTH_INCHES,
        first_column_inches=CHEMISTRY_LABEL_COLUMN_WIDTH_INCHES,
    )


def _add_properties_table(document: Document, *, detail: QuartaTaglioDetailResponse) -> None:
    _add_bullet_section_title(document, "Mechanical properties: raw material acc. to EN 755-2")
    available_properties = {
        item.field: item
        for item in detail.properties
        if item.value is not None or item.standard_min is not None or item.standard_max is not None
    }
    if not available_properties:
        document.add_paragraph("No mechanical properties available.")
        return
    fields = [field for field in PROPERTY_WORD_FIELDS if field in available_properties or field in {"diametro", "S"}]
    extra_fields = [field for field in available_properties if field not in fields]
    fields.extend(extra_fields)
    table = _new_table(document, [""] + [_property_header(field) for field in fields])
    for cell in table.rows[0].cells:
        _set_cell_no_wrap(cell)
    min_row = table.add_row().cells
    spec_row = table.add_row().cells
    min_row[0].text = "Min."
    spec_row[0].text = "Spec."
    for index, field in enumerate(fields, start=1):
        item = available_properties.get(field)
        min_row[index].text = _format_value(item.standard_min if item else None)
        spec_row[index].text = _property_spec_value(item, field=field)
    _set_table_font(table, size=10)
    _set_table_cell_margins(table, left=35, right=35)
    for row in table.rows:
        for cell in row.cells:
            _set_cell_no_wrap(cell)
    _bold_row(spec_row)
    _set_table_width(table, Inches(7.1))


def _add_notes(document: Document, *, detail: QuartaTaglioDetailResponse) -> None:
    ok_notes = [item for item in detail.notes if item.status == "ok"]
    if not ok_notes:
        return
    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(0)
    run = title.add_run("Notes:")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    for item in ok_notes:
        text = item.value
        if text:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1
            run = paragraph.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)


def _add_footer_signatures(document: Document, *, certified_by: User, quality_manager: User | None) -> None:
    section = document.sections[0]
    section.different_first_page_header_footer = False
    _fill_signature_footer(section.footer, certified_by=certified_by, quality_manager=quality_manager, include_operator=True)


def _fill_signature_footer(footer, *, certified_by: User, quality_manager: User | None, include_operator: bool) -> None:
    footer.is_linked_to_previous = False
    _clear_header_footer(footer)

    table = footer.add_table(rows=1, cols=3, width=Inches(6.7))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _clear_table_borders(table)
    _set_column_widths(table, [Inches(3.4), Inches(1.8), Inches(1.5)])
    for cell in table.rows[0].cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    operator_paragraph = table.cell(0, 0).paragraphs[0]
    operator_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    operator_paragraph.paragraph_format.space_after = Pt(0)
    if include_operator:
        _add_signature_label_value(operator_paragraph, "Operator:", certified_by.name)

    qm_paragraph = table.cell(0, 1).paragraphs[0]
    qm_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    qm_paragraph.paragraph_format.space_after = Pt(0)
    _add_signature_label_value(qm_paragraph, "Quality Manager:", "")
    signature_paragraph = table.cell(0, 2).paragraphs[0]
    signature_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    signature_paragraph.paragraph_format.space_after = Pt(0)
    if QUALITY_MANAGER_SIGNATURE_PATH.exists():
        signature_paragraph.add_run().add_picture(str(QUALITY_MANAGER_SIGNATURE_PATH), width=Inches(0.75))
    elif quality_manager:
        signature_paragraph.add_run(quality_manager.name)


def _append_docx_body(
    base_docx_path: Path,
    extra_docx_path: Path,
    *,
    detail: QuartaTaglioDetailResponse,
    draft_number: str,
    certified_by: User,
    quality_manager: User | None,
) -> None:
    """Append the uploaded extra-pages DOCX as real Word body content.

    ``docxcompose`` handles the Word relationship graph (images, styles and
    numbering) much more safely than a manual XML merge. This avoids the Microsoft
    Word "recovered unreadable content" warning seen with altChunk/manual merges.
    """
    master = Document(str(base_docx_path))
    extra_document = Document(str(extra_docx_path))
    master.add_page_break()
    composer = Composer(master)
    composer.append(extra_document)
    composer.save(str(base_docx_path))
    merged = Document(str(base_docx_path))
    _apply_document_shell(
        merged,
        detail=detail,
        draft_number=draft_number,
        certified_by=certified_by,
        quality_manager=quality_manager,
    )
    merged.save(str(base_docx_path))


def _append_pdf_attachments(
    base_docx_path: Path,
    pdf_attachments: Sequence[tuple[Path, str]],
    *,
    certified_by: User,
    quality_manager: User | None,
) -> None:
    document = Document(str(base_docx_path))
    with tempfile.TemporaryDirectory(prefix="certi_pdf_attachment_") as tmp_root:
        tmp_dir = Path(tmp_root)
        attachment_section = None
        for pdf_path, original_filename in pdf_attachments:
            if not pdf_path.exists():
                continue
            pdf_document = fitz.open(str(pdf_path))
            for page_index, page in enumerate(pdf_document, start=1):
                if attachment_section is None:
                    attachment_section = document.add_section(WD_SECTION.NEW_PAGE)
                    attachment_section.top_margin = Inches(0.45)
                    attachment_section.bottom_margin = Inches(0.65)
                    attachment_section.left_margin = Inches(0.45)
                    attachment_section.right_margin = Inches(0.45)
                    attachment_section.header_distance = Inches(0.12)
                    attachment_section.footer_distance = Inches(0.12)
                    attachment_section.different_first_page_header_footer = False
                    attachment_section.header.is_linked_to_previous = False
                    attachment_section.footer.is_linked_to_previous = False
                    _clear_header_footer(attachment_section.header)
                    _fill_signature_footer(attachment_section.footer, certified_by=certified_by, quality_manager=quality_manager, include_operator=True)
                else:
                    document.add_page_break()

                pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                image_path = tmp_dir / f"{_safe_attachment_name(original_filename)}_{page_index}.png"
                pixmap.save(image_path)
                _add_pdf_attachment_page(document, image_path=image_path, original_filename=original_filename, section=attachment_section)
            pdf_document.close()
    document.save(str(base_docx_path))


def _add_pdf_attachment_page(document: Document, *, image_path: Path, original_filename: str, section) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run(f"Attachment: {original_filename}")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(9)

    with Image.open(image_path) as image:
        width_px, height_px = image.size
    if width_px <= 0 or height_px <= 0:
        return

    usable_width = max(1.0, (int(section.page_width) - int(section.left_margin) - int(section.right_margin)) / 914400)
    usable_height = max(
        1.0,
        (int(section.page_height) - int(section.top_margin) - int(section.bottom_margin)) / 914400 - 0.75,
    )
    aspect_ratio = width_px / height_px
    target_width = usable_width
    target_height = target_width / aspect_ratio
    if target_height > usable_height:
        target_height = usable_height
        target_width = target_height * aspect_ratio

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_picture(str(image_path), width=Inches(target_width), height=Inches(target_height))


def _safe_attachment_name(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9_-]+", "_", value).strip("_")[:80] or "attachment"


def _clear_header_footer(container) -> None:
    for child in list(container._element):
        container._element.remove(child)


def _add_field(paragraph, instruction: str) -> None:
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), instruction)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    field.append(run)
    paragraph._p.append(field)


def _add_content_control_run(
    paragraph,
    *,
    tag: str,
    text: object,
    font_name: str,
    size: int,
    bold: bool = False,
    color: str | None = None,
) -> None:
    sdt = OxmlElement("w:sdt")
    sdt_pr = OxmlElement("w:sdtPr")
    control_id = OxmlElement("w:id")
    control_id.set(qn("w:val"), str(_content_control_id(tag)))
    alias = OxmlElement("w:alias")
    alias.set(qn("w:val"), tag)
    tag_element = OxmlElement("w:tag")
    tag_element.set(qn("w:val"), tag)
    text_control = OxmlElement("w:text")
    sdt_pr.append(control_id)
    sdt_pr.append(alias)
    sdt_pr.append(tag_element)
    sdt_pr.append(text_control)
    sdt.append(sdt_pr)

    content = OxmlElement("w:sdtContent")
    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), font_name)
    fonts.set(qn("w:hAnsi"), font_name)
    run_pr.append(fonts)
    if bold:
        run_pr.append(OxmlElement("w:b"))
    if color:
        color_element = OxmlElement("w:color")
        color_element.set(qn("w:val"), color)
        run_pr.append(color_element)
    size_element = OxmlElement("w:sz")
    size_element.set(qn("w:val"), str(size * 2))
    run_pr.append(size_element)
    run.append(run_pr)
    text_element = OxmlElement("w:t")
    text_element.set(qn("xml:space"), "preserve")
    text_element.text = "" if text in (None, "") else str(text)
    run.append(text_element)
    content.append(run)
    sdt.append(content)
    paragraph._p.append(sdt)


def _content_control_id(tag: str) -> int:
    return zlib.crc32(f"certi_nt:{tag}".encode("utf-8")) & 0x7FFFFFFF


def inspect_docx_content_controls(path: Path) -> tuple[list[str], list[str]]:
    present: set[str] = set()
    with zipfile.ZipFile(path) as archive:
        for name in archive.namelist():
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            xml = archive.read(name).decode("utf-8", errors="ignore")
            for tag in CONTENT_CONTROL_TAGS:
                if f'w:val="{tag}"' in xml or f"w:val='{tag}'" in xml:
                    present.add(tag)
    ordered_present = [tag for tag in CONTENT_CONTROL_TAGS if tag in present]
    missing = [tag for tag in CONTENT_CONTROL_TAGS if tag not in present]
    return ordered_present, missing


def update_docx_content_controls(source_path: Path, output_path: Path, values: dict[str, object]) -> tuple[list[str], list[str]]:
    tmp_path = output_path.with_suffix(f"{output_path.suffix}.tmp")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    updated_tags: set[str] = set()
    with zipfile.ZipFile(source_path, "r") as source, zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            content = source.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                text = content.decode("utf-8", errors="ignore")
                for tag, value in values.items():
                    if tag not in CONTENT_CONTROL_TAGS:
                        continue
                    text, changed = _replace_content_control_text(text, tag, "" if value in (None, "") else str(value))
                    if changed:
                        updated_tags.add(tag)
                content = text.encode("utf-8")
            target.writestr(item, content)
    tmp_path.replace(output_path)
    normalize_certificate_docx_layout(output_path)
    present, missing = inspect_docx_content_controls(output_path)
    return [tag for tag in CONTENT_CONTROL_TAGS if tag in updated_tags], missing


def normalize_certificate_docx_layout(path: Path) -> None:
    """Apply the exact Word/PDF-safe header values selected in the audit tests."""
    tmp_path = path.with_suffix(f"{path.suffix}.layout.tmp")
    changed = False
    try:
        with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as target:
            for item in source.infolist():
                content = source.read(item.filename)
                if item.filename.startswith("word/header") and item.filename.endswith(".xml"):
                    xml = content.decode("utf-8", errors="ignore")
                    normalized = _normalize_certificate_header_xml(xml)
                    if normalized != xml:
                        content = normalized.encode("utf-8")
                        changed = True
                target.writestr(item, content)
    except (OSError, zipfile.BadZipFile):
        tmp_path.unlink(missing_ok=True)
        return
    if changed:
        tmp_path.replace(path)
    else:
        tmp_path.unlink(missing_ok=True)


def _normalize_certificate_header_xml(xml: str) -> str:
    if not _header_xml_looks_like_certificate(xml):
        return xml
    xml = _normalize_logo_xml(xml)
    xml = _normalize_header_flow_xml(xml)
    return xml


def _header_xml_looks_like_certificate(xml: str) -> bool:
    return (
        "EN 10204 - 3.1" in xml
        or "Certificate n" in xml
        or 'w:val="CERT_NUMBER"' in xml
        or 'w:val="COD_F3_RAW"' in xml
    )


def _normalize_logo_xml(xml: str) -> str:
    xml = re.sub(
        r'(<wp:extent\b[^>]*\bcx=")\d+("[^>]*\bcy=")\d+("[^>]*/?>)',
        rf"\g<1>{HEADER_LOGO_WIDTH_EMU}\g<2>{HEADER_LOGO_HEIGHT_EMU}\g<3>",
        xml,
        count=1,
    )
    xml = re.sub(
        r'(<a:ext\b[^>]*\bcx=")\d+("[^>]*\bcy=")\d+("[^>]*/?>)',
        rf"\g<1>{HEADER_LOGO_WIDTH_EMU}\g<2>{HEADER_LOGO_HEIGHT_EMU}\g<3>",
        xml,
        count=1,
    )
    return re.sub(
        r'<w:spacing\s+w:before="0"\s+w:after="0"\s+w:line="192"\s+w:lineRule="auto"\s*/>',
        '<w:spacing w:before="0" w:after="0"/>',
        xml,
        count=1,
    )


def _normalize_header_flow_xml(xml: str) -> str:
    if 'w:val="COD_F3_RAW"' not in xml and "Cod. F3 Raw:" not in xml:
        return xml
    xml = re.sub(
        r'<w:tblCellMar><w:left w:w="\d+" w:type="dxa"\s*/><w:right w:w="\d+" w:type="dxa"\s*/><w:top w:w="\d+" w:type="dxa"\s*/><w:bottom w:w="\d+" w:type="dxa"\s*/></w:tblCellMar>',
        '<w:tblCellMar><w:left w:w="26" w:type="dxa"/><w:right w:w="26" w:type="dxa"/><w:top w:w="7" w:type="dxa"/><w:bottom w:w="7" w:type="dxa"/></w:tblCellMar>',
        xml,
        count=1,
    )
    xml = re.sub(
        r'<w:trHeight w:val="\d+" w:hRule="atLeast"\s*/>',
        '<w:trHeight w:val="395" w:hRule="atLeast"/>',
        xml,
        count=4,
    )
    return re.sub(
        r'(<w:spacing w:after="0" w:line=")\d+(" w:lineRule="auto"\s*/>)',
        r'\g<1>220\g<2>',
        xml,
    )


def _replace_content_control_text(xml: str, tag: str, value: str) -> tuple[str, bool]:
    tag_marker = f'w:val="{tag}"'
    start = 0
    changed = False
    while True:
        tag_index = xml.find(tag_marker, start)
        if tag_index < 0:
            break
        sdt_start = _find_content_control_start(xml, tag_index)
        sdt_end = xml.find("</w:sdt>", tag_index)
        if sdt_start < 0 or sdt_end < 0:
            start = tag_index + len(tag_marker)
            continue
        sdt_end += len("</w:sdt>")
        block = xml[sdt_start:sdt_end]
        content_start = block.find("<w:sdtContent")
        if content_start < 0:
            start = sdt_end
            continue
        text_start = _find_word_text_start(block, content_start)
        if text_start < 0:
            start = sdt_end
            continue
        text_open_end = block.find(">", text_start)
        text_close = block.find("</w:t>", text_open_end)
        if text_open_end < 0:
            start = sdt_end
            continue
        escaped = _escape_xml_text(value)
        if text_close < 0:
            if text_open_end > 0 and block[text_open_end - 1] == "/":
                block = block[: text_open_end - 1] + ">" + escaped + "</w:t>" + block[text_open_end + 1 :]
            else:
                start = sdt_end
                continue
        else:
            block = block[: text_open_end + 1] + escaped + block[text_close:]
        xml = xml[:sdt_start] + block + xml[sdt_end:]
        changed = True
        start = sdt_start + len(block)
    return xml, changed


def _find_content_control_start(xml: str, before_index: int) -> int:
    exact = xml.rfind("<w:sdt>", 0, before_index)
    with_attrs = xml.rfind("<w:sdt ", 0, before_index)
    return max(exact, with_attrs)


def _find_word_text_start(xml: str, from_index: int) -> int:
    exact = xml.find("<w:t>", from_index)
    with_attrs = xml.find("<w:t ", from_index)
    candidates = [index for index in (exact, with_attrs) if index >= 0]
    return min(candidates) if candidates else -1


def _escape_xml_text(value: str) -> str:
    return value.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _force_word_field_update(document: Document) -> None:
    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def _remove_page_number_restart(section) -> None:
    sect_pr = section._sectPr
    for page_num_type in list(sect_pr.findall(qn("w:pgNumType"))):
        sect_pr.remove(page_num_type)


def _add_section_title(document: Document, text: str) -> None:
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = 0.4
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(51, 51, 51)


def _add_bullet_section_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(22)
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run(f"•  {text}")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)


def _add_box_table(document: Document, rows: list[tuple[tuple[str, object], tuple[str, object]]]) -> None:
    table = document.add_table(rows=0, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left[0]
        cells[1].text = _empty_dash(left[1])
        cells[2].text = right[0]
        cells[3].text = _empty_dash(right[1])
        for cell_index in (0, 2):
            for paragraph in cells[cell_index].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
    _set_table_font(table, size=8)


def _add_header_data_table(container, rows: list[tuple[tuple[str, str, object], tuple[str, str, object]]]) -> None:
    try:
        table = container.add_table(rows=0, cols=2, width=Inches(7.1))
    except TypeError:
        table = container.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for left, right in rows:
        cells = table.add_row().cells
        _fill_header_data_cell(cells[0], left[0], left[1], left[2])
        _fill_header_data_cell(cells[1], right[0], right[1], right[2])
    _set_table_width(table, Inches(7.1))


def _add_certificate_header_flow_table(container, certificate_header: dict[str, object]) -> None:
    table = container.add_table(rows=0, cols=3, width=HEADER_WIDTH)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_cell_margins(table, left=26, right=26, top=7, bottom=7)
    _set_column_widths(table, HEADER_FLOW_COLUMN_WIDTHS)
    rows = [
        (
            ("Purchaser:", "Cliente:", certificate_header.get("cliente"), "PURCHASER"),
            ("Cod. F3 Raw:", "Cod. F3 del grezzo:", certificate_header.get("codice_f3_raw"), "COD_F3_RAW"),
            ("Cod. F3 Finished:", "Cod. F3 del finito:", certificate_header.get("codice_f3_finished"), "COD_F3_FINISHED"),
        ),
        (
            ("Order.:", "Ordine:", certificate_header.get("ordine_cliente"), "ORDER_CLIENT"),
            ("Drawing / Description Raw:", "Disegno / Descrizione del grezzo:", certificate_header.get("descrizione_raw"), "RAW_DESCRIPTION"),
            ("Drawing / Description Finished:", "Disegno / Descrizione del finito:", certificate_header.get("descrizione_finished"), "FINISHED_DESCRIPTION"),
        ),
        (
            ("Confirm of order:", "C.d.O.:", certificate_header.get("conferma_ordine"), "CONFIRM_ORDER"),
            ("D.d.T.:", "D.d.T.:", certificate_header.get("ddt_raw"), "DDT_RAW"),
            ("D.d.T.:", "D.d.T.:", certificate_header.get("ddt_finished"), "DDT_FINISHED"),
        ),
        (
            ("", "", "", None),
            ("Quantity:", "Quantità:", certificate_header.get("quantita_raw"), "QUANTITY_RAW"),
            ("Quantity:", "Quantità:", certificate_header.get("quantita_finished"), "QUANTITY_FINISHED"),
        ),
    ]
    for cells_data in rows:
        row = table.add_row()
        row.height = HEADER_FLOW_ROW_HEIGHT
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        cells = row.cells
        for cell, (english_label, italian_label, value, control_tag) in zip(cells, cells_data):
            _fill_header_block_cell(cell, english_label, italian_label, value, control_tag=control_tag)
    _set_table_width(table, HEADER_WIDTH)
    _set_column_widths(table, HEADER_FLOW_COLUMN_WIDTHS)


def _fill_header_block_cell(cell, english_label: str, italian_label: str, value: object, *, control_tag: str | None = None) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = HEADER_FLOW_LINE_SPACING
    if english_label:
        label_run = paragraph.add_run(english_label)
        label_run.bold = True
        label_run.font.name = "Arial"
        label_run.font.size = Pt(10)
    if italian_label:
        italian = cell.add_paragraph()
        italian.alignment = WD_ALIGN_PARAGRAPH.CENTER
        italian.paragraph_format.space_after = Pt(0)
        italian.paragraph_format.line_spacing = HEADER_FLOW_LINE_SPACING
        run = italian.add_run(italian_label)
        run.font.name = "Arial"
        run.font.size = Pt(6.5)
    value_text = "" if value in (None, "") else str(value)
    if value_text or control_tag:
        value_paragraph = cell.add_paragraph()
        value_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        value_paragraph.paragraph_format.space_after = Pt(0)
        value_paragraph.paragraph_format.line_spacing = HEADER_FLOW_LINE_SPACING
        if control_tag:
            _add_content_control_run(
                value_paragraph,
                tag=control_tag,
                text=value_text,
                font_name="Times New Roman",
                size=9 if len(value_text) > 42 else 10,
                color="0070C0",
            )
        else:
            value_run = value_paragraph.add_run(value_text)
            value_run.font.name = "Times New Roman"
            value_run.font.size = Pt(9 if len(value_text) > 42 else 10)
            value_run.font.color.rgb = RGBColor(0, 112, 192)


def _fill_header_data_cell(cell, english_label: str, italian_label: str, value: object) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    if english_label:
        label_run = paragraph.add_run(f"{english_label} ")
        label_run.bold = True
        label_run.font.name = "Arial"
        label_run.font.size = Pt(12)
    if value not in (None, ""):
        value_run = paragraph.add_run(str(value))
        value_run.font.name = "Times New Roman"
        value_run.font.size = Pt(12)
    italian = cell.add_paragraph()
    italian.alignment = WD_ALIGN_PARAGRAPH.LEFT
    italian.paragraph_format.space_after = Pt(0)
    italian.paragraph_format.line_spacing = 1
    run = italian.add_run(italian_label)
    run.font.name = "Arial"
    run.font.size = Pt(8)


def _add_inline_label_value(paragraph, label: str, value: object, *, control_tag: str | None = None) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_run = paragraph.add_run(f"{label} ")
    label_run.bold = True
    label_run.font.name = "Times New Roman"
    label_run.font.size = Pt(12)
    if control_tag:
        _add_content_control_run(
            paragraph,
            tag=control_tag,
            text=_empty_dash(value),
            font_name="Times New Roman",
            size=11 if len(_empty_dash(value)) > 8 else 12,
        )
        return
    value_run = paragraph.add_run(_empty_dash(value))
    value_run.font.name = "Times New Roman"
    value_run.font.size = Pt(12)


def _add_signature_label_value(paragraph, label: str, value: object) -> None:
    label_run = paragraph.add_run(f"{label} ")
    label_run.bold = True
    label_run.font.name = "Times New Roman"
    label_run.font.size = Pt(12)
    value_text = _empty_dash(value) if value not in (None, "") else ""
    if value_text:
        value_run = paragraph.add_run(value_text)
        value_run.font.name = "Times New Roman"
        value_run.font.size = Pt(12)


def _alloy_label(detail: QuartaTaglioDetailResponse) -> str:
    if detail.selected_standard and detail.selected_standard.label:
        parts = [part.strip() for part in detail.selected_standard.label.split("·") if part.strip()]
        return " ".join(parts[:3]) if parts else detail.selected_standard.label
    return "-"


def _alloy_code(detail: QuartaTaglioDetailResponse) -> str:
    if detail.selected_standard and detail.selected_standard.label:
        return detail.selected_standard.label.split("·")[0].strip()
    return ""


def _property_header(field: str) -> str:
    return PROPERTY_HEADER_LABELS.get(field, field)


def _property_spec_value(item, *, field: str) -> str:
    if field in {"diametro", "S", "Ø", "Diametro"} and item is None:
        return "/"
    return _format_value(item.value if item else None)


def _new_table(document: Document, headers: list[str]):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
        for paragraph in header_cells[index].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    _set_table_font(table, size=8)
    return table


def _set_table_font(table, *, size: int) -> None:
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(size)


def _bold_row(cells) -> None:
    for cell in cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def _set_table_width(table, width) -> None:
    table.autofit = False
    table.allow_autofit = False
    column_count = len(table.columns)
    if not column_count:
        return
    column_width = width / column_count
    for column in table.columns:
        for cell in column.cells:
            cell.width = column_width


def _set_column_widths(table, widths) -> None:
    table.autofit = False
    table.allow_autofit = False
    for column, width in zip(table.columns, widths):
        for cell in column.cells:
            cell.width = width


def _set_table_width_with_first_column(table, *, total_inches: float, first_column_inches: float) -> None:
    table.autofit = False
    table.allow_autofit = False
    column_count = len(table.columns)
    if column_count == 0:
        return
    if column_count == 1:
        _set_column_widths(table, [Inches(total_inches)])
        return
    remaining_inches = max(total_inches - first_column_inches, 0.1)
    other_column_inches = remaining_inches / (column_count - 1)
    widths = [Inches(first_column_inches)] + [Inches(other_column_inches)] * (column_count - 1)
    _set_column_widths(table, widths)


def _set_table_cell_margins(table, *, left: int, right: int, top: int | None = None, bottom: int | None = None) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    values = {"left": left, "right": right}
    if top is not None:
        values["top"] = top
    if bottom is not None:
        values["bottom"] = bottom
    for edge, value in values.items():
        element = tbl_cell_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tbl_cell_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_cell_no_wrap(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    if tc_pr.find(qn("w:noWrap")) is None:
        tc_pr.append(OxmlElement("w:noWrap"))


def _clear_table_borders(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.first_child_found_in("w:tcBorders")
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                tag = f"w:{edge}"
                element = borders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    borders.append(element)
                element.set(qn("w:val"), "nil")

def _empty_dash(value: object) -> str:
    if value is None or value == "":
        return "-"
    return str(value)


def _format_value(value: float | str | None) -> str:
    if value is None or value == "":
        return "-"
    if isinstance(value, float):
        return f"{value:.2f}".rstrip("0").rstrip(".").replace(".", ",")
    return str(value)


def _format_quantity(value: float | str | None) -> str:
    if value is None or value == "":
        return "-"
    if isinstance(value, str):
        normalized = value.strip().replace(",", ".")
        try:
            value = float(normalized)
        except ValueError:
            return value.replace(".", ",")
    if isinstance(value, (int, float)):
        return str(int(round(value)))
    return str(value)


def _format_chemistry_value(value: float | str | None) -> str:
    if value is None or value == "":
        return "-"
    if isinstance(value, float):
        return f"{value:.3f}".rstrip("0").rstrip(".").replace(".", ",")
    return str(value)
